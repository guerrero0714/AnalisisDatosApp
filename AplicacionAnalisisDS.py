import sys
import pandas as pd
import numpy as np
from PyQt5.QtWidgets import (QApplication, QMainWindow, QTabWidget, QStatusBar, 
                            QAction, QToolBar, QVBoxLayout, QWidget, QLabel, 
                            QTableView, QFileDialog, QPushButton, QComboBox, 
                            QGroupBox, QHBoxLayout, QScrollArea, QCheckBox,
                            QLineEdit, QSpinBox, QDoubleSpinBox, QMessageBox,
                            QSplitter, QFormLayout, QListWidget, QListWidgetItem,
                            QTextEdit, QDialog)
from PyQt5.QtCore import Qt, QAbstractTableModel, QRegExp
from PyQt5.QtGui import QIcon, QRegExpValidator
from matplotlib.backends.backend_qt5agg import FigureCanvasQTAgg as FigureCanvas
from matplotlib.figure import Figure
import seaborn as sns
from scipy import stats
import warnings
from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
from io import BytesIO
from docx2pdf import convert
from htmldocx import HtmlToDocx

warnings.filterwarnings('ignore')

class DataFrameModel(QAbstractTableModel):
    """Modelo avanzado para mostrar DataFrames de pandas en QTableView"""
    def __init__(self, data=pd.DataFrame()):
        QAbstractTableModel.__init__(self)
        self._data = data
        self._filtered_data = data
        self._filters = {}
    
    def rowCount(self, parent=None):
        return self._filtered_data.shape[0]
    
    def columnCount(self, parent=None):
        return self._filtered_data.shape[1]
    
    def data(self, index, role=Qt.DisplayRole):
        if index.isValid():
            if role == Qt.DisplayRole:
                value = self._filtered_data.iloc[index.row(), index.column()]
                if pd.isna(value):
                    return ""
                return str(value)
            elif role == Qt.TextAlignmentRole:
                return Qt.AlignCenter
        return None
    
    def headerData(self, section, orientation, role):
        if role == Qt.DisplayRole:
            if orientation == Qt.Horizontal:
                return str(self._filtered_data.columns[section])
            if orientation == Qt.Vertical:
                return str(self._filtered_data.index[section])
        return None
    
    def set_dataframe(self, dataframe):
        self.beginResetModel()
        self._data = dataframe.copy()
        self._filtered_data = dataframe.copy()
        self._filters = {}
        self.endResetModel()
    
    def apply_filters(self, filters):
        """Aplica múltiples filtros a los datos"""
        self.beginResetModel()
        self._filters = filters
        self._filtered_data = self._data.copy()
        
        for col, (op, value) in filters.items():
            if op == "contains":
                self._filtered_data = self._filtered_data[self._filtered_data[col].astype(str).str.contains(value, case=False, na=False)]
            elif op == "==":
                self._filtered_data = self._filtered_data[self._filtered_data[col] == value]
            elif op == "!=":
                self._filtered_data = self._filtered_data[self._filtered_data[col] != value]
            elif op == ">":
                self._filtered_data = self._filtered_data[self._filtered_data[col] > value]
            elif op == "<":
                self._filtered_data = self._filtered_data[self._filtered_data[col] < value]
            elif op == ">=":
                self._filtered_data = self._filtered_data[self._filtered_data[col] >= value]
            elif op == "<=":
                self._filtered_data = self._filtered_data[self._filtered_data[col] <= value]
            elif op == "is_null":
                self._filtered_data = self._filtered_data[self._filtered_data[col].isna()]
            elif op == "not_null":
                self._filtered_data = self._filtered_data[self._filtered_data[col].notna()]
        
        self.endResetModel()
        return self._filtered_data

class AdvancedDataVisualizer(QWidget):
    """Widget avanzado para visualización de gráficos con múltiples opciones"""
    def __init__(self):
        super().__init__()
        self.layout = QVBoxLayout(self)
        self.figure = Figure(figsize=(10, 6), dpi=100)
        self.canvas = FigureCanvas(self.figure)
        self.toolbar = self.create_toolbar()
        
        self.layout.addWidget(self.toolbar)
        self.layout.addWidget(self.canvas)
        
        self.current_plot_type = None
        self.current_data = None
    
    def create_toolbar(self):
        """Crea una barra de herramientas para los gráficos"""
        toolbar = QToolBar("Herramientas de gráfico")
        
        save_action = QAction(QIcon.fromTheme("document-save"), "Guardar gráfico", self)
        save_action.triggered.connect(self.save_plot)
        toolbar.addAction(save_action)
        
        zoom_action = QAction(QIcon.fromTheme("zoom-in"), "Zoom", self)
        zoom_action.triggered.connect(self.toggle_zoom)
        toolbar.addAction(zoom_action)
        
        return toolbar
    
    def save_plot(self):
        """Guarda el gráfico actual en un archivo"""
        if not hasattr(self, 'current_plot_type') or not self.current_plot_type:
            return
            
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar gráfico",
            "",
            "PNG (*.png);;JPEG (*.jpg);;PDF (*.pdf);;SVG (*.svg)",
            options=options
        )
        
        if file_path:
            self.figure.savefig(file_path, dpi=300, bbox_inches='tight')
    
    def toggle_zoom(self):
        """Activa/desactiva el zoom (implementación básica)"""
        pass
    
    def plot_histogram(self, data, column, bins=20, kde=True, rug=False):
        """Genera un histograma avanzado"""
        self.current_plot_type = "histogram"
        self.current_data = (data, column)
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        sns.histplot(data=data, x=column, bins=bins, kde=kde, ax=ax)
        if rug:
            sns.rugplot(data=data, x=column, ax=ax)
        
        ax.set_title(f'Distribución de {column}')
        self.canvas.draw()
    
    def plot_scatter(self, data, x_col, y_col, hue=None, size=None, style=None):
        """Genera gráfico de dispersión avanzado"""
        self.current_plot_type = "scatter"
        self.current_data = (data, x_col, y_col)
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        sns.scatterplot(data=data, x=x_col, y=y_col, hue=hue, size=size, style=style, ax=ax)
        ax.set_title(f'{y_col} vs {x_col}')
        self.canvas.draw()
    
    def plot_bar(self, data, x_col, y_col, hue=None, ci=None):
        """Genera gráfico de barras avanzado"""
        self.current_plot_type = "bar"
        self.current_data = (data, x_col, y_col)
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        sns.barplot(data=data, x=x_col, y=y_col, hue=hue, ci=ci, ax=ax)
        ax.set_title(f'{y_col} por {x_col}')
        self.canvas.draw()
    
    def plot_box(self, data, x_col, y_col, hue=None):
        """Genera diagrama de caja avanzado"""
        self.current_plot_type = "box"
        self.current_data = (data, x_col, y_col)
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        sns.boxplot(data=data, x=x_col, y=y_col, hue=hue, ax=ax)
        ax.set_title(f'Distribución de {y_col} por {x_col}')
        self.canvas.draw()
    
    def plot_correlation_matrix(self, data):
        """Genera matriz de correlación"""
        self.current_plot_type = "correlation"
        self.current_data = data
        self.figure.clear()
        ax = self.figure.add_subplot(111)
        
        numeric_data = data.select_dtypes(include=[np.number])
        corr_matrix = numeric_data.corr()
        sns.heatmap(corr_matrix, annot=True, fmt=".2f", cmap='coolwarm', center=0, ax=ax)
        ax.set_title('Matriz de Correlación')
        self.canvas.draw()

class FilterWidget(QGroupBox):
    """Widget avanzado para aplicar filtros a los datos"""
    def __init__(self, parent=None):
        super().__init__("Filtros Avanzados")
        self.parent = parent
        self.layout = QVBoxLayout(self)
        
        # Selector de columnas
        self.column_combo = QComboBox()
        self.column_combo.currentIndexChanged.connect(self.update_filter_options)
        
        # Operadores de filtro
        self.operator_combo = QComboBox()
        
        # Valor del filtro
        self.value_input = QLineEdit()
        self.value_input.setPlaceholderText("Valor a filtrar")
        
        # Botones
        self.add_filter_btn = QPushButton("Añadir Filtro")
        self.add_filter_btn.clicked.connect(self.add_filter)
        self.clear_filters_btn = QPushButton("Limpiar Filtros")
        self.clear_filters_btn.clicked.connect(self.clear_filters)
        
        # Lista de filtros activos
        self.active_filters_list = QListWidget()
        
        # Diseño
        form_layout = QFormLayout()
        form_layout.addRow("Columna:", self.column_combo)
        form_layout.addRow("Operador:", self.operator_combo)
        form_layout.addRow("Valor:", self.value_input)
        
        self.layout.addLayout(form_layout)
        self.layout.addWidget(self.add_filter_btn)
        self.layout.addWidget(self.clear_filters_btn)
        self.layout.addWidget(QLabel("Filtros Activos:"))
        self.layout.addWidget(self.active_filters_list)
    
    def set_columns(self, columns):
        """Establece las columnas disponibles para filtrar"""
        self.column_combo.clear()
        self.column_combo.addItems(columns)
        self.update_filter_options()
    
    def update_filter_options(self):
        """Actualiza los operadores disponibles según el tipo de columna"""
        current_col = self.column_combo.currentText()
        if not current_col or not hasattr(self.parent, 'current_data'):
            return
            
        col_type = self.parent.current_data[current_col].dtype
        
        self.operator_combo.clear()
        
        # Operadores comunes para todos los tipos
        operators = ["==", "!=", "is_null", "not_null"]
        
        # Operadores para texto
        if col_type == 'object':
            operators.extend(["contains", "starts_with", "ends_with"])
        
        # Operadores para números
        if np.issubdtype(col_type, np.number):
            operators.extend([">", "<", ">=", "<=", "between"])
        
        self.operator_combo.addItems(operators)
    
    def add_filter(self):
        """Añade un nuevo filtro"""
        col = self.column_combo.currentText()
        op = self.operator_combo.currentText()
        value = self.value_input.text()
        
        if not col:
            return
            
        # Validar valor para operadores numéricos
        if op in [">", "<", ">=", "<=", "==", "!="] and not value:
            QMessageBox.warning(self, "Error", "Por favor ingrese un valor para el filtro")
            return
        
        # Crear representación del filtro para mostrar en la lista
        filter_text = f"{col} {op}"
        if op not in ["is_null", "not_null"]:
            filter_text += f" {value}"
        
        item = QListWidgetItem(filter_text)
        item.setData(Qt.UserRole, (col, op, value))
        self.active_filters_list.addItem(item)
        
        # Aplicar filtros
        self.apply_filters()
    
    def clear_filters(self):
        """Limpia todos los filtros"""
        self.active_filters_list.clear()
        self.apply_filters()
    
    def apply_filters(self):
        """Aplica todos los filtros activos a los datos"""
        if not hasattr(self.parent, 'data_model'):
            return
            
        filters = {}
        
        for i in range(self.active_filters_list.count()):
            item = self.active_filters_list.item(i)
            col, op, value = item.data(Qt.UserRole)
            
            # Convertir valor para operadores numéricos
            if op in [">", "<", ">=", "<=", "==", "!="] and self.parent.current_data[col].dtype != 'object':
                try:
                    if '.' in value:
                        value = float(value)
                    else:
                        value = int(value)
                except ValueError:
                    continue
            
            filters[col] = (op, value)
        
        self.parent.data_model.apply_filters(filters)
        self.parent.update_stats()

class StatsTestsWidget(QGroupBox):
    """Widget para realizar pruebas estadísticas"""
    def __init__(self, parent=None):
        super().__init__("Pruebas Estadísticas")
        self.parent = parent
        self.layout = QVBoxLayout(self)
        
        # Selector de prueba
        self.test_combo = QComboBox()
        self.test_combo.addItems([
            "Test t para muestras independientes",
            "ANOVA unidireccional",
            "Correlación de Pearson",
            "Test de normalidad (Shapiro-Wilk)"
        ])
        self.test_combo.currentIndexChanged.connect(self.update_test_ui)
        
        # Variables para la prueba
        self.var1_combo = QComboBox()
        self.var2_combo = QComboBox()
        self.group_combo = QComboBox()
        
        # Resultados
        self.results_text = QLabel("Seleccione una prueba y las variables")
        self.results_text.setWordWrap(True)
        
        # Botón para ejecutar
        self.run_test_btn = QPushButton("Ejecutar Prueba")
        self.run_test_btn.clicked.connect(self.run_test)
        
        # Diseño inicial
        self.test_params_layout = QFormLayout()
        
        self.layout.addWidget(QLabel("Prueba:"))
        self.layout.addWidget(self.test_combo)
        self.layout.addLayout(self.test_params_layout)
        self.layout.addWidget(self.run_test_btn)
        self.layout.addWidget(self.results_text)
        self.layout.addStretch()
    
    def update_columns(self, columns):
        """Actualiza las columnas disponibles para análisis"""
        numeric_cols = [col for col in columns if self.parent.current_data[col].dtype.kind in 'fbiu']
        cat_cols = [col for col in columns if self.parent.current_data[col].dtype == 'object']
        
        self.var1_combo.clear()
        self.var1_combo.addItems(numeric_cols)
        
        self.var2_combo.clear()
        self.var2_combo.addItems(numeric_cols)
        
        self.group_combo.clear()
        self.group_combo.addItems(cat_cols)
    
    def update_test_ui(self):
        """Actualiza la UI según la prueba seleccionada"""
        # Limpiar layout de parámetros
        while self.test_params_layout.count():
            child = self.test_params_layout.takeAt(0)
            if child.widget():
                child.widget().deleteLater()
        
        test_name = self.test_combo.currentText()
        
        if test_name == "Test t para muestras independientes":
            self.test_params_layout.addRow("Variable numérica:", self.var1_combo)
            self.test_params_layout.addRow("Variable categórica (2 grupos):", self.group_combo)
        elif test_name == "ANOVA unidireccional":
            self.test_params_layout.addRow("Variable numérica:", self.var1_combo)
            self.test_params_layout.addRow("Variable categórica (3+ grupos):", self.group_combo)
        elif test_name == "Correlación de Pearson":
            self.test_params_layout.addRow("Variable 1:", self.var1_combo)
            self.test_params_layout.addRow("Variable 2:", self.var2_combo)
        elif test_name == "Test de normalidad (Shapiro-Wilk)":
            self.test_params_layout.addRow("Variable numérica:", self.var1_combo)
    
    def run_test(self):
        """Ejecuta la prueba estadística seleccionada"""
        if not hasattr(self.parent, 'current_data') or self.parent.current_data is None:
            self.results_text.setText("No hay datos cargados")
            return
            
        test_name = self.test_combo.currentText()
        data = self.parent.current_data
        results = ""
        
        try:
            if test_name == "Test t para muestras independientes":
                var = self.var1_combo.currentText()
                group = self.group_combo.currentText()
                
                # Verificar que hay exactamente 2 grupos
                groups = data[group].unique()
                if len(groups) != 2:
                    self.results_text.setText("Error: La variable categórica debe tener exactamente 2 grupos")
                    return
                
                group1 = data[data[group] == groups[0]][var]
                group2 = data[data[group] == groups[1]][var]
                
                t_stat, p_value = stats.ttest_ind(group1, group2, nan_policy='omit')
                
                results = f"Test t para muestras independientes\n\n"
                results += f"Variable: {var}\n"
                results += f"Grupos: {groups[0]} vs {groups[1]}\n\n"
                results += f"Estadístico t: {t_stat:.4f}\n"
                results += f"Valor p: {p_value:.4f}\n\n"
                results += "Interpretación:\n"
                if p_value < 0.05:
                    results += "Hay una diferencia estadísticamente significativa entre los grupos (p < 0.05)"
                else:
                    results += "No hay evidencia de diferencia significativa entre los grupos (p >= 0.05)"
            
            elif test_name == "ANOVA unidireccional":
                var = self.var1_combo.currentText()
                group = self.group_combo.currentText()
                
                groups = data[group].unique()
                if len(groups) < 2:
                    self.results_text.setText("Error: La variable categórica debe tener al menos 2 grupos")
                    return
                
                # Preparar datos para ANOVA
                group_data = [data[data[group] == g][var].dropna() for g in groups]
                
                f_stat, p_value = stats.f_oneway(*group_data)
                
                results = f"ANOVA unidireccional\n\n"
                results += f"Variable: {var}\n"
                results += f"Grupos: {', '.join(map(str, groups))}\n\n"
                results += f"Estadístico F: {f_stat:.4f}\n"
                results += f"Valor p: {p_value:.4f}\n\n"
                results += "Interpretación:\n"
                if p_value < 0.05:
                    results += "Hay diferencias estadísticamente significativas entre al menos dos grupos (p < 0.05)"
                else:
                    results += "No hay evidencia de diferencias significativas entre los grupos (p >= 0.05)"
            
            elif test_name == "Correlación de Pearson":
                var1 = self.var1_combo.currentText()
                var2 = self.var2_combo.currentText()
                
                clean_data = data[[var1, var2]].dropna()
                
                corr, p_value = stats.pearsonr(clean_data[var1], clean_data[var2])
                
                results = f"Correlación de Pearson\n\n"
                results += f"Variables: {var1} y {var2}\n\n"
                results += f"Coeficiente de correlación: {corr:.4f}\n"
                results += f"Valor p: {p_value:.4f}\n\n"
                results += "Interpretación:\n"
                if p_value < 0.05:
                    if abs(corr) > 0.7:
                        strength = "fuerte"
                    elif abs(corr) > 0.3:
                        strength = "moderada"
                    else:
                        strength = "débil"
                    
                    direction = "positiva" if corr > 0 else "negativa"
                    results += f"Hay una correlación {strength} y {direction} estadísticamente significativa (p < 0.05)"
                else:
                    results += "No hay evidencia de correlación significativa (p >= 0.05)"
            
            elif test_name == "Test de normalidad (Shapiro-Wilk)":
                var = self.var1_combo.currentText()
                
                test_data = data[var].dropna()
                
                if len(test_data) < 3:
                    self.results_text.setText("Error: Se necesitan al menos 3 observaciones para la prueba")
                    return
                
                stat, p_value = stats.shapiro(test_data)
                
                results = f"Test de normalidad Shapiro-Wilk\n\n"
                results += f"Variable: {var}\n\n"
                results += f"Estadístico W: {stat:.4f}\n"
                results += f"Valor p: {p_value:.4f}\n\n"
                results += "Interpretación:\n"
                if p_value < 0.05:
                    results += "Los datos NO siguen una distribución normal (p < 0.05)"
                else:
                    results += "Los datos podrían seguir una distribución normal (p >= 0.05)"
            
            self.results_text.setText(results)
        
        except Exception as e:
            self.results_text.setText(f"Error al ejecutar la prueba: {str(e)}")

class ReportGenerator(QGroupBox):
    """Widget para generar informes personalizados"""
    def __init__(self, parent=None):
        super().__init__("Generador de Informes")
        self.parent = parent
        self.layout = QVBoxLayout(self)
        
        # Configuración del informe
        self.report_title = QLineEdit()
        self.report_title.setPlaceholderText("Título del informe")
        
        self.report_author = QLineEdit()
        self.report_author.setPlaceholderText("Autor del informe")
        
        self.include_stats = QCheckBox("Incluir estadísticas descriptivas")
        self.include_stats.setChecked(True)
        
        self.include_correlations = QCheckBox("Incluir matriz de correlación")
        self.include_correlations.setChecked(True)
        
        self.include_pivot = QCheckBox("Incluir tabla dinámica seleccionada")
        
        self.include_chart = QCheckBox("Incluir gráfico seleccionado")
        
        # Selección de secciones personalizadas
        self.sections_list = QListWidget()
        self.sections_list.setSelectionMode(QListWidget.MultiSelection)
        
        # Botones
        self.generate_btn = QPushButton("Generar Informe")
        self.generate_btn.clicked.connect(self.generate_report)
        
        self.preview_btn = QPushButton("Vista Previa")
        self.preview_btn.clicked.connect(self.preview_report)
        
        # Diseño
        form_layout = QFormLayout()
        form_layout.addRow("Título:", self.report_title)
        form_layout.addRow("Autor:", self.report_author)
        
        self.layout.addLayout(form_layout)
        self.layout.addWidget(self.include_stats)
        self.layout.addWidget(self.include_correlations)
        self.layout.addWidget(self.include_pivot)
        self.layout.addWidget(self.include_chart)
        self.layout.addWidget(QLabel("Secciones adicionales:"))
        self.layout.addWidget(self.sections_list)
        
        btn_layout = QHBoxLayout()
        btn_layout.addWidget(self.generate_btn)
        btn_layout.addWidget(self.preview_btn)
        self.layout.addLayout(btn_layout)
    
    def update_sections(self):
        """Actualiza las secciones disponibles basadas en los datos"""
        self.sections_list.clear()
        
        if not hasattr(self.parent, 'current_data') or self.parent.current_data is None:
            return
            
        # Secciones basadas en columnas
        for col in self.parent.current_data.columns:
            item = QListWidgetItem(f"Análisis de {col}")
            item.setData(Qt.UserRole, ("column", col))
            self.sections_list.addItem(item)
        
        # Secciones basadas en análisis especiales
        special_analyses = [
            ("Distribución de valores faltantes", "missing_values"),
            ("Análisis de valores atípicos", "outliers"),
            ("Resumen por grupos", "group_summary")
        ]
        
        for name, analysis_type in special_analyses:
            item = QListWidgetItem(name)
            item.setData(Qt.UserRole, ("analysis", analysis_type))
            self.sections_list.addItem(item)
    
    def generate_report(self):
        """Genera un informe completo"""
        if not hasattr(self.parent, 'current_data') or self.parent.current_data is None:
            QMessageBox.warning(self, "Error", "No hay datos cargados para generar el informe")
            return
            
        # Obtener opciones del informe
        title = self.report_title.text() or "Informe de Análisis"
        author = self.report_author.text() or ""
        
        # Obtener rutas para guardar
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Guardar informe",
            "",
            "Documento Word (*.docx);;PDF (*.pdf);;HTML (*.html)",
            options=options
        )
        
        if not file_path:
            return
            
        try:
            # Crear documento Word
            doc = Document()
            
            # Título y metadatos
            doc.add_heading(title, level=0)
            if author:
                doc.add_paragraph(f"Autor: {author}")
            doc.add_paragraph(f"Fecha: {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}")
            doc.add_paragraph("\n")
            
            # Resumen ejecutivo
            doc.add_heading("Resumen Ejecutivo", level=1)
            doc.add_paragraph(self.get_executive_summary())
            
            # Estadísticas descriptivas
            if self.include_stats.isChecked():
                doc.add_heading("Estadísticas Descriptivas", level=1)
                doc.add_paragraph("A continuación se presentan las estadísticas descriptivas básicas para cada variable:")
                
                # Agregar tabla de estadísticas
                stats = self.parent.current_data.describe(include='all').T
                stats['missing'] = self.parent.current_data.isna().sum()
                stats['unique'] = self.parent.current_data.nunique()
                
                table = doc.add_table(stats.shape[0]+1, stats.shape[1])
                
                # Encabezados
                for j, col in enumerate(stats.columns):
                    table.cell(0, j).text = str(col)
                
                # Datos
                for i, index in enumerate(stats.index):
                    table.cell(i+1, 0).text = str(index)
                    for j, col in enumerate(stats.columns):
                        table.cell(i+1, j+1).text = str(stats.loc[index, col])
                
                doc.add_paragraph("\n")
            
            # Matriz de correlación
            if self.include_correlations.isChecked():
                numeric_cols = self.parent.current_data.select_dtypes(include=[np.number]).columns
                if len(numeric_cols) > 1:
                    doc.add_heading("Matriz de Correlación", level=1)
                    
                    # Generar y guardar gráfico temporalmente
                    fig, ax = plt.subplots(figsize=(10, 8))
                    sns.heatmap(
                        self.parent.current_data[numeric_cols].corr(), 
                        annot=True, fmt=".2f", cmap='coolwarm', center=0, ax=ax
                    )
                    plt.title('Matriz de Correlación')
                    
                    memfile = BytesIO()
                    fig.savefig(memfile, format='png', dpi=300, bbox_inches='tight')
                    doc.add_picture(memfile, width=Inches(6))
                    plt.close(fig)
                    
                    doc.add_paragraph("\n")
            
            # Tabla dinámica
            if self.include_pivot.isChecked() and hasattr(self.parent, 'pivot_model'):
                doc.add_heading("Tabla Dinámica", level=1)
                
                # Agregar tabla al documento
                pivot_data = self.parent.pivot_model._filtered_data
                table = doc.add_table(pivot_data.shape[0]+1, pivot_data.shape[1])
                
                # Encabezados
                if isinstance(pivot_data.columns, pd.MultiIndex):
                    for j, col in enumerate(pivot_data.columns):
                        table.cell(0, j).text = " | ".join(map(str, col))
                else:
                    for j, col in enumerate(pivot_data.columns):
                        table.cell(0, j).text = str(col)
                
                # Datos
                for i, index in enumerate(pivot_data.index):
                    table.cell(i+1, 0).text = str(index)
                    for j, col in enumerate(pivot_data.columns):
                        table.cell(i+1, j+1).text = str(pivot_data.loc[index, col])
                
                doc.add_paragraph("\n")
            
            # Gráfico seleccionado
            if self.include_chart.isChecked() and hasattr(self.parent, 'visualizer'):
                doc.add_heading("Gráfico", level=1)
                
                # Guardar gráfico actual en memoria
                fig = self.parent.visualizer.figure
                if fig and len(fig.axes) > 0:
                    memfile = BytesIO()
                    fig.savefig(memfile, format='png', dpi=300, bbox_inches='tight')
                    doc.add_picture(memfile, width=Inches(6))
                    doc.add_paragraph("\n")
            
            # Secciones seleccionadas
            selected_items = self.sections_list.selectedItems()
            for item in selected_items:
                section_type, value = item.data(Qt.UserRole)
                
                if section_type == "column":
                    doc.add_heading(f"Análisis de {value}", level=1)
                    doc.add_paragraph(self.get_column_analysis(value))
                
                elif section_type == "analysis" and value == "missing_values":
                    doc.add_heading("Distribución de Valores Faltantes", level=1)
                    doc.add_paragraph(self.get_missing_values_analysis())
                
                elif section_type == "analysis" and value == "outliers":
                    doc.add_heading("Análisis de Valores Atípicos", level=1)
                    doc.add_paragraph(self.get_outliers_analysis())
                
                elif section_type == "analysis" and value == "group_summary":
                    doc.add_heading("Resumen por Grupos", level=1)
                    doc.add_paragraph(self.get_group_summary())
            
            # Guardar documento
            if file_path.endswith('.docx'):
                doc.save(file_path)
            elif file_path.endswith('.pdf'):
                # Convertir DOCX a PDF
                convert(file_path)
            elif file_path.endswith('.html'):
                # Implementar exportación a HTML
                self.export_to_html(doc, file_path)
            
            self.parent.status_bar.showMessage(f"Informe generado en {file_path}", 5000)
            QMessageBox.information(self, "Informe generado", f"El informe se ha guardado en:\n{file_path}")
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo generar el informe:\n{str(e)}")
    
    def preview_report(self):
        """Muestra una vista previa del informe"""
        preview_dialog = QDialog(self)
        preview_dialog.setWindowTitle("Vista Previa del Informe")
        preview_dialog.setMinimumSize(800, 600)
        
        layout = QVBoxLayout(preview_dialog)
        text_edit = QTextEdit()
        text_edit.setReadOnly(True)
        
        # Generar contenido del informe
        report_content = []
        report_content.append(f"# {self.report_title.text() or 'Informe de Análisis'}\n")
        report_content.append(f"**Autor:** {self.report_author.text() or 'Anónimo'}\n")
        report_content.append(f"**Fecha:** {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}\n")
        report_content.append("\n## Resumen Ejecutivo\n")
        report_content.append(self.get_executive_summary())
        
        if self.include_stats.isChecked():
            report_content.append("\n## Estadísticas Descriptivas\n")
            report_content.append("Estadísticas descriptivas básicas para cada variable:\n")
            
            stats = self.parent.current_data.describe(include='all').T
            stats['missing'] = self.parent.current_data.isna().sum()
            stats['unique'] = self.parent.current_data.nunique()
            
            # Formatear como tabla Markdown
            report_content.append("| Variable | " + " | ".join(stats.columns) + " |")
            report_content.append("|----------|" + "|".join(["----------"] * len(stats.columns)) + "|")
            
            for index, row in stats.iterrows():
                report_content.append(f"| {index} | " + " | ".join([str(round(x, 4)) if isinstance(x, (int, float)) else str(x) for x in row]) + " |")
        
        # Mostrar en el diálogo
        text_edit.setMarkdown("\n".join(report_content))
        layout.addWidget(text_edit)
        
        preview_dialog.exec_()
    
    def get_executive_summary(self):
        """Genera un resumen ejecutivo de los datos"""
        data = self.parent.current_data
        num_cols = data.select_dtypes(include=[np.number]).columns
        cat_cols = data.select_dtypes(include=['object', 'category']).columns
        
        summary = []
        summary.append(f"El conjunto de datos contiene {len(data)} filas y {len(data.columns)} columnas.")
        summary.append(f"Incluye {len(num_cols)} variables numéricas y {len(cat_cols)} variables categóricas.")
        
        missing_total = data.isna().sum().sum()
        if missing_total > 0:
            summary.append(f"Se encontraron {missing_total} valores faltantes en total.")
        
        dup_total = data.duplicated().sum()
        if dup_total > 0:
            summary.append(f"Existen {dup_total} filas duplicadas en los datos.")
        
        if len(num_cols) > 1:
            corr_matrix = data[num_cols].corr().abs()
            high_corr = corr_matrix.unstack().sort_values(ascending=False).drop_duplicates()
            high_corr = high_corr[(high_corr > 0.7) & (high_corr < 1.0)]
            
            if len(high_corr) > 0:
                summary.append("\n**Correlaciones fuertes encontradas:**")
                for pair, corr in high_corr.items():
                    summary.append(f"- {pair[0]} y {pair[1]}: {corr:.2f}")
        
        return "\n".join(summary)
    
    def get_column_analysis(self, column):
        """Genera un análisis detallado para una columna específica"""
        data = self.parent.current_data[column]
        analysis = []
        
        if pd.api.types.is_numeric_dtype(data):
            analysis.append(f"**Variable numérica:** {column}")
            analysis.append(f"- Media: {data.mean():.2f}")
            analysis.append(f"- Mediana: {data.median():.2f}")
            analysis.append(f"- Desviación estándar: {data.std():.2f}")
            analysis.append(f"- Rango: [{data.min():.2f}, {data.max():.2f}]")
            analysis.append(f"- Asimetría: {data.skew():.2f}")
            analysis.append(f"- Curtosis: {data.kurtosis():.2f}")
            
            # Análisis de outliers usando IQR
            q1 = data.quantile(0.25)
            q3 = data.quantile(0.75)
            iqr = q3 - q1
            lower_bound = q1 - 1.5 * iqr
            upper_bound = q3 + 1.5 * iqr
            
            outliers = data[(data < lower_bound) | (data > upper_bound)]
            if len(outliers) > 0:
                analysis.append(f"- **Valores atípicos:** {len(outliers)} ({len(outliers)/len(data)*100:.1f}%)")
        
        else:
            analysis.append(f"**Variable categórica:** {column}")
            analysis.append(f"- Valores únicos: {data.nunique()}")
            analysis.append(f"- Valor más frecuente: {data.mode()[0]} (aparece {data.value_counts().max()} veces)")
            
            if data.nunique() < 10:
                analysis.append("\n**Distribución de valores:**")
                for value, count in data.value_counts().items():
                    analysis.append(f"- {value}: {count} ({count/len(data)*100:.1f}%)")
        
        missing = data.isna().sum()
        if missing > 0:
            analysis.append(f"\n**Valores faltantes:** {missing} ({missing/len(data)*100:.1f}%)")
        
        return "\n".join(analysis)
    
    def get_missing_values_analysis(self):
        """Analiza la distribución de valores faltantes"""
        data = self.parent.current_data
        missing = data.isna().sum()
        missing = missing[missing > 0]
        
        if len(missing) == 0:
            return "No hay valores faltantes en el conjunto de datos."
        
        analysis = []
        analysis.append("**Distribución de valores faltantes por columna:**")
        
        for col, count in missing.items():
            analysis.append(f"- {col}: {count} valores faltantes ({count/len(data)*100:.1f}%)")
        
        # Patrones de valores faltantes
        analysis.append("\n**Patrones de valores faltantes:**")
        
        # Columnas que faltan juntas
        missing_matrix = data.isna()
        correlations = missing_matrix.corr()
        
        high_corr = correlations.unstack().sort_values(ascending=False).drop_duplicates()
        high_corr = high_corr[(high_corr > 0.5) & (high_corr < 1.0)]
        
        if len(high_corr) > 0:
            analysis.append("Las siguientes columnas tienden a tener valores faltantes juntos:")
            for pair, corr in high_corr.items():
                analysis.append(f"- {pair[0]} y {pair[1]}: correlación {corr:.2f}")
        else:
            analysis.append("No se encontraron patrones fuertes de valores faltantes entre columnas.")
        
        return "\n".join(analysis)
    
    def get_outliers_analysis(self):
        """Identifica y analiza valores atípicos"""
        data = self.parent.current_data
        numeric_cols = data.select_dtypes(include=[np.number]).columns
        
        if len(numeric_cols) == 0:
            return "No hay columnas numéricas para analizar valores atípicos."
        
        analysis = []
        analysis.append("**Resumen de valores atípicos por columna numérica:**\n")
        
        for col in numeric_cols:
            col_data = data[col].dropna()
            q1 = col_data.quantile(0.25)
            q3 = col_data.quantile(0.75)
            iqr = q3 - q1
            
            lower_bound = q1 - 1.5 * iqr
            upper_bound = q3 + 1.5 * iqr
            
            outliers = col_data[(col_data < lower_bound) | (col_data > upper_bound)]
            
            if len(outliers) > 0:
                analysis.append(f"**{col}**:")
                analysis.append(f"- Valores atípicos: {len(outliers)} ({len(outliers)/len(col_data)*100:.1f}%)")
                analysis.append(f"- Rango de atípicos: [{outliers.min():.2f}, {outliers.max():.2f}]")
                analysis.append(f"- Valores extremos (>3*IQR): {len(col_data[(col_data < q1 - 3*iqr) | (col_data > q3 + 3*iqr)])}")
                analysis.append("")
        
        if len(analysis) == 1:  # Solo el título
            return "No se encontraron valores atípicos significativos en las columnas numéricas."
        
        return "\n".join(analysis)
    
    def get_group_summary(self):
        """Genera un resumen por grupos para columnas categóricas"""
        data = self.parent.current_data
        cat_cols = data.select_dtypes(include=['object', 'category']).columns
        
        if len(cat_cols) == 0:
            return "No hay columnas categóricas para agrupar."
        
        analysis = []
        analysis.append("**Resumen por grupos:**\n")
        
        for col in cat_cols:
            if data[col].nunique() > 20:
                analysis.append(f"La columna '{col}' tiene demasiados valores únicos ({data[col].nunique()}) para un análisis de grupo útil.")
                continue
            
            analysis.append(f"**Agrupado por: {col}**")
            
            # Para cada grupo, mostrar estadísticas de columnas numéricas
            numeric_cols = data.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) > 0:
                group_stats = data.groupby(col)[numeric_cols].agg(['mean', 'std', 'count'])
                
                for num_col in numeric_cols:
                    analysis.append(f"\n*{num_col}* por grupo:")
                    for group in group_stats.index:
                        mean = group_stats.loc[group, (num_col, 'mean')]
                        std = group_stats.loc[group, (num_col, 'std')]
                        count = group_stats.loc[group, (num_col, 'count')]
                        analysis.append(f"- {group}: Media={mean:.2f}, DE={std:.2f} (n={count})")
            
            analysis.append("\n")
        
        return "\n".join(analysis)
    
    def export_to_html(self, doc, file_path):
        """Exporta el informe a formato HTML (implementación básica)"""
        html_content = []
        html_content.append("<!DOCTYPE html>")
        html_content.append("<html>")
        html_content.append("<head>")
        html_content.append("<meta charset='UTF-8'>")
        html_content.append(f"<title>{self.report_title.text() or 'Informe de Análisis'}</title>")
        html_content.append("<style>")
        html_content.append("body { font-family: Arial, sans-serif; line-height: 1.6; margin: 0 auto; max-width: 900px; padding: 20px; }")
        html_content.append("h1 { color: #2c3e50; border-bottom: 2px solid #3498db; padding-bottom: 10px; }")
        html_content.append("h2 { color: #2980b9; border-bottom: 1px solid #3498db; padding-bottom: 5px; }")
        html_content.append("table { border-collapse: collapse; width: 100%; margin: 20px 0; }")
        html_content.append("th, td { border: 1px solid #ddd; padding: 8px; text-align: left; }")
        html_content.append("th { background-color: #f2f2f2; }")
        html_content.append("tr:nth-child(even) { background-color: #f9f9f9; }")
        html_content.append("</style>")
        html_content.append("</head>")
        html_content.append("<body>")
        
        # Título
        html_content.append(f"<h1>{self.report_title.text() or 'Informe de Análisis'}</h1>")
        html_content.append(f"<p><strong>Autor:</strong> {self.report_author.text() or 'Anónimo'}</p>")
        html_content.append(f"<p><strong>Fecha:</strong> {pd.Timestamp.now().strftime('%Y-%m-%d %H:%M')}</p>")
        
        # Resumen ejecutivo
        html_content.append("<h2>Resumen Ejecutivo</h2>")
        html_content.append(f"<p>{self.get_executive_summary().replace('\n', '<br>')}</p>")
        
        # Estadísticas descriptivas
        if self.include_stats.isChecked():
            html_content.append("<h2>Estadísticas Descriptivas</h2>")
            html_content.append("<p>A continuación se presentan las estadísticas descriptivas básicas para cada variable:</p>")
            
            stats = self.parent.current_data.describe(include='all').T
            stats['missing'] = self.parent.current_data.isna().sum()
            stats['unique'] = self.parent.current_data.nunique()
            
            # Tabla HTML
            html_content.append("<table>")
            html_content.append("<tr><th>Variable</th>" + "".join(f"<th>{col}</th>" for col in stats.columns) + "</tr>")
            
            for index, row in stats.iterrows():
                html_content.append("<tr>")
                html_content.append(f"<td>{index}</td>")
                for col in stats.columns:
                    html_content.append(f"<td>{str(round(row[col], 4)) if isinstance(row[col], (int, float)) else str(row[col])}</td>")
                html_content.append("</tr>")
            
            html_content.append("</table>")
        
        # Secciones adicionales
        selected_items = self.sections_list.selectedItems()
        for item in selected_items:
            section_type, value = item.data(Qt.UserRole)
            
            if section_type == "column":
                html_content.append(f"<h2>Análisis de {value}</h2>")
                html_content.append(f"<p>{self.get_column_analysis(value).replace('\n', '<br>')}</p>")
            
            elif section_type == "analysis" and value == "missing_values":
                html_content.append("<h2>Distribución de Valores Faltantes</h2>")
                html_content.append(f"<p>{self.get_missing_values_analysis().replace('\n', '<br>')}</p>")
            
            elif section_type == "analysis" and value == "outliers":
                html_content.append("<h2>Análisis de Valores Atípicos</h2>")
                html_content.append(f"<p>{self.get_outliers_analysis().replace('\n', '<br>')}</p>")
            
            elif section_type == "analysis" and value == "group_summary":
                html_content.append("<h2>Resumen por Grupos</h2>")
                html_content.append(f"<p>{self.get_group_summary().replace('\n', '<br>')}</p>")
        
        html_content.append("</body>")
        html_content.append("</html>")
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write("\n".join(html_content))

class DataAnalyzerPro(QMainWindow):
    """Aplicación avanzada de análisis de datos con generador de informes"""
    def __init__(self):
        super().__init__()
        self.current_data = None
        self.init_ui()
        self.create_actions()
        self.create_menus()
        self.create_toolbar()
        
    def init_ui(self):
        """Inicializa la interfaz de usuario avanzada"""
        self.setWindowTitle("DataAnalyzer Pro - Análisis Avanzado de Datos")
        self.setGeometry(100, 100, 1400, 900)
        
        # Widget central con splitter
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        self.main_layout = QHBoxLayout(central_widget)
        
        # Splitter principal
        self.main_splitter = QSplitter(Qt.Horizontal)
        self.main_layout.addWidget(self.main_splitter)
        
        # Panel izquierdo (datos y filtros)
        left_panel = QWidget()
        left_layout = QVBoxLayout(left_panel)
        
        # Panel derecho (análisis y gráficos)
        right_panel = QWidget()
        right_layout = QVBoxLayout(right_panel)
        
        # Configuración del panel izquierdo
        self.setup_data_tab(left_layout)
        self.setup_filter_panel(left_layout)
        
        # Configuración del panel derecho
        self.setup_analysis_tabs(right_layout)
        
        # Añadir paneles al splitter
        self.main_splitter.addWidget(left_panel)
        self.main_splitter.addWidget(right_panel)
        self.main_splitter.setSizes([400, 1000])
        
        # Barra de estado
        self.status_bar = QStatusBar()
        self.setStatusBar(self.status_bar)
    
    def setup_data_tab(self, layout):
        """Configura el panel de visualización de datos"""
        data_group = QGroupBox("Datos")
        data_layout = QVBoxLayout(data_group)
        
        # Tabla de datos
        self.data_table = QTableView()
        self.data_table.setSelectionBehavior(QTableView.SelectRows)
        self.data_table.setSelectionMode(QTableView.ExtendedSelection)
        self.data_model = DataFrameModel()
        self.data_table.setModel(self.data_model)
        
        # Barra de herramientas de datos
        data_toolbar = QToolBar("Herramientas de datos")
        self.export_data_action = QAction(QIcon.fromTheme("document-save"), "Exportar datos", self)
        self.export_data_action.triggered.connect(self.export_data)
        data_toolbar.addAction(self.export_data_action)
        
        data_layout.addWidget(data_toolbar)
        data_layout.addWidget(self.data_table)
        
        layout.addWidget(data_group)
    
    def setup_filter_panel(self, layout):
        """Configura el panel de filtros avanzados"""
        self.filter_widget = FilterWidget(self)
        layout.addWidget(self.filter_widget)
    
    def setup_analysis_tabs(self, layout):
        """Configura las pestañas de análisis avanzado"""
        self.analysis_tabs = QTabWidget()
        
        # Pestaña de estadísticas descriptivas
        self.setup_stats_tab()
        
        # Pestaña de tablas dinámicas
        self.setup_pivot_tab()
        
        # Pestaña de gráficos avanzados
        self.setup_charts_tab()
        
        # Pestaña de pruebas estadísticas
        self.setup_stat_tests_tab()
        
        # Pestaña de informes
        self.setup_reports_tab()
        
        layout.addWidget(self.analysis_tabs)
    
    def setup_stats_tab(self):
        """Configura la pestaña de estadísticas descriptivas avanzadas"""
        stats_tab = QWidget()
        stats_layout = QVBoxLayout(stats_tab)
        
        # Estadísticas descriptivas
        stats_group = QGroupBox("Estadísticas Descriptivas Avanzadas")
        stats_group_layout = QVBoxLayout(stats_group)
        
        self.stats_table = QTableView()
        self.stats_model = DataFrameModel()
        self.stats_table.setModel(self.stats_model)
        
        stats_group_layout.addWidget(self.stats_table)
        stats_layout.addWidget(stats_group)
        
        # Estadísticas por grupo
        group_stats_group = QGroupBox("Estadísticas por Grupo")
        group_stats_layout = QVBoxLayout(group_stats_group)
        
        self.group_combo = QComboBox()
        self.group_combo.currentIndexChanged.connect(self.update_group_stats)
        
        self.group_stats_table = QTableView()
        self.group_stats_model = DataFrameModel()
        self.group_stats_table.setModel(self.group_stats_model)
        
        group_stats_layout.addWidget(QLabel("Agrupar por:"))
        group_stats_layout.addWidget(self.group_combo)
        group_stats_layout.addWidget(self.group_stats_table)
        stats_layout.addWidget(group_stats_group)
        
        stats_layout.addStretch()
        self.analysis_tabs.addTab(stats_tab, "Estadísticas")
    
    def setup_pivot_tab(self):
        """Configura la pestaña de tablas dinámicas avanzadas"""
        pivot_tab = QWidget()
        layout = QVBoxLayout(pivot_tab)
        
        # Constructor de tablas dinámicas
        pivot_group = QGroupBox("Tablas Dinámicas Avanzadas")
        pivot_layout = QVBoxLayout(pivot_group)
        
        # Controles para tabla dinámica
        controls_layout = QHBoxLayout()
        
        self.pivot_row_combo = QComboBox()
        self.pivot_col_combo = QComboBox()
        self.pivot_val_combo = QComboBox()
        self.pivot_agg_combo = QComboBox()
        self.pivot_agg_combo.addItems(["count", "sum", "mean", "median", "min", "max", "std", "var"])
        
        controls_layout.addWidget(QLabel("Filas:"))
        controls_layout.addWidget(self.pivot_row_combo)
        controls_layout.addWidget(QLabel("Columnas:"))
        controls_layout.addWidget(self.pivot_col_combo)
        controls_layout.addWidget(QLabel("Valores:"))
        controls_layout.addWidget(self.pivot_val_combo)
        controls_layout.addWidget(QLabel("Agregación:"))
        controls_layout.addWidget(self.pivot_agg_combo)
        
        pivot_layout.addLayout(controls_layout)
        
        # Botones adicionales
        btn_layout = QHBoxLayout()
        
        self.pivot_build_btn = QPushButton("Generar Tabla Dinámica")
        self.pivot_build_btn.clicked.connect(self.build_advanced_pivot)
        
        self.pivot_export_btn = QPushButton("Exportar Tabla")
        self.pivot_export_btn.clicked.connect(self.export_pivot)
        
        btn_layout.addWidget(self.pivot_build_btn)
        btn_layout.addWidget(self.pivot_export_btn)
        pivot_layout.addLayout(btn_layout)
        
        # Resultados
        self.pivot_result = QTableView()
        self.pivot_model = DataFrameModel()
        self.pivot_result.setModel(self.pivot_model)
        pivot_layout.addWidget(self.pivot_result)
        
        layout.addWidget(pivot_group)
        self.analysis_tabs.addTab(pivot_tab, "Tablas Dinámicas")
    
    def setup_charts_tab(self):
        """Configura la pestaña de gráficos avanzados"""
        charts_tab = QWidget()
        layout = QVBoxLayout(charts_tab)
        
        # Controles para gráficos
        chart_controls = QHBoxLayout()
        
        self.chart_type_combo = QComboBox()
        self.chart_type_combo.addItems([
            "Histograma", 
            "Dispersión", 
            "Barras", 
            "Caja", 
            "Matriz de Correlación",
            "Gráfico de Violín",
            "Gráfico de Densidad"
        ])
        
        self.chart_x_combo = QComboBox()
        self.chart_y_combo = QComboBox()
        self.chart_hue_combo = QComboBox()
        self.chart_hue_combo.addItem("Ninguno", "")
        
        chart_controls.addWidget(QLabel("Tipo:"))
        chart_controls.addWidget(self.chart_type_combo)
        chart_controls.addWidget(QLabel("Eje X:"))
        chart_controls.addWidget(self.chart_x_combo)
        chart_controls.addWidget(QLabel("Eje Y:"))
        chart_controls.addWidget(self.chart_y_combo)
        chart_controls.addWidget(QLabel("Color por:"))
        chart_controls.addWidget(self.chart_hue_combo)
        
        # Opciones adicionales según tipo de gráfico
        self.chart_options_group = QGroupBox("Opciones del Gráfico")
        self.chart_options_layout = QHBoxLayout(self.chart_options_group)
        
        # Botón para generar
        self.plot_btn = QPushButton("Generar Gráfico")
        self.plot_btn.clicked.connect(self.generate_advanced_chart)
        
        # Visualizador de gráficos
        self.visualizer = AdvancedDataVisualizer()
        
        layout.addLayout(chart_controls)
        layout.addWidget(self.chart_options_group)
        layout.addWidget(self.plot_btn)
        layout.addWidget(self.visualizer)
        
        self.analysis_tabs.addTab(charts_tab, "Gráficos")
    
    def setup_stat_tests_tab(self):
        """Configura la pestaña de pruebas estadísticas"""
        tests_tab = QWidget()
        layout = QVBoxLayout(tests_tab)
        
        self.stats_tests_widget = StatsTestsWidget(self)
        layout.addWidget(self.stats_tests_widget)
        
        self.analysis_tabs.addTab(tests_tab, "Pruebas Estadísticas")
    
    def setup_reports_tab(self):
        """Configura la pestaña de generación de informes"""
        reports_tab = QWidget()
        reports_layout = QVBoxLayout(reports_tab)
        
        self.report_generator = ReportGenerator(self)
        reports_layout.addWidget(self.report_generator)
        
        self.analysis_tabs.addTab(reports_tab, "Informes")
    
    def create_actions(self):
        """Crea las acciones avanzadas del menú"""
        # Acción para abrir archivo
        self.open_action = QAction(QIcon.fromTheme("document-open"), "Abrir archivo", self)
        self.open_action.setShortcut("Ctrl+O")
        self.open_action.setStatusTip("Abrir un archivo de datos")
        self.open_action.triggered.connect(self.load_file)
        
        # Acción para exportar
        self.export_action = QAction(QIcon.fromTheme("document-save-as"), "Exportar", self)
        self.export_action.setShortcut("Ctrl+E")
        self.export_action.setStatusTip("Exportar datos o resultados")
        self.export_action.triggered.connect(self.export_data)
        
        # Acción para salir
        self.exit_action = QAction(QIcon.fromTheme("application-exit"), "Salir", self)
        self.exit_action.setShortcut("Ctrl+Q")
        self.exit_action.setStatusTip("Salir de la aplicación")
        self.exit_action.triggered.connect(self.close)
        
        # Acción para análisis rápido
        self.quick_analysis_action = QAction(QIcon.fromTheme("document-properties"), "Análisis Rápido", self)
        self.quick_analysis_action.setStatusTip("Realizar un análisis rápido de los datos")
        self.quick_analysis_action.triggered.connect(self.run_quick_analysis)
    
    def create_menus(self):
        """Crea los menús avanzados de la aplicación"""
        menubar = self.menuBar()
        
        # Menú Archivo
        file_menu = menubar.addMenu("Archivo")
        file_menu.addAction(self.open_action)
        file_menu.addAction(self.export_action)
        file_menu.addSeparator()
        file_menu.addAction(self.exit_action)
        
        # Menú Análisis
        analysis_menu = menubar.addMenu("Análisis")
        analysis_menu.addAction(self.quick_analysis_action)
        
        # Menú Ayuda
        help_menu = menubar.addMenu("Ayuda")
        about_action = QAction("Acerca de", self)
        about_action.triggered.connect(self.show_about)
        help_menu.addAction(about_action)
    
    def create_toolbar(self):
        """Crea la barra de herramientas avanzada"""
        toolbar = QToolBar("Barra de herramientas principal")
        self.addToolBar(toolbar)
        
        toolbar.addAction(self.open_action)
        toolbar.addAction(self.export_action)
        toolbar.addSeparator()
        toolbar.addAction(self.quick_analysis_action)
    
    def load_file(self):
        """Carga un archivo de datos con más opciones"""
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getOpenFileName(
            self,
            "Cargar archivo de datos",
            "",
            "Archivos de datos (*.xlsx *.csv *.txt);;Excel (*.xlsx);;CSV (*.csv);;Texto (*.txt);;Todos los archivos (*)",
            options=options
        )
        
        if not file_path:
            return
        
        try:
            # Opciones de carga para CSV
            if file_path.endswith('.csv') or file_path.endswith('.txt'):
                # Diálogo para opciones de CSV
                delimiter = "," if file_path.endswith('.csv') else "\t"
                header = "infer"  # Por defecto intenta detectar encabezados
                
                self.current_data = pd.read_csv(file_path, delimiter=delimiter, header=header)
            
            elif file_path.endswith('.xlsx'):
                # Para Excel, podríamos añadir opción de seleccionar hoja
                self.current_data = pd.read_excel(file_path)
            
            else:
                raise ValueError("Formato de archivo no soportado")
            
            self.update_ui_with_data()
            self.status_bar.showMessage("Datos cargados correctamente", 3000)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo cargar el archivo:\n{str(e)}")
            self.status_bar.showMessage(f"Error: {str(e)}", 5000)
    
    def update_ui_with_data(self):
        """Actualiza toda la UI con los datos cargados"""
        if self.current_data is not None:
            # Actualizar tabla de datos principal
            self.data_model.set_dataframe(self.current_data)
            
            # Actualizar estadísticas
            self.update_stats()
            
            # Actualizar filtros
            self.filter_widget.set_columns(self.current_data.columns.tolist())
            
            # Actualizar combos de tablas dinámicas
            self.update_pivot_controls()
            
            # Actualizar controles de gráficos
            self.update_chart_controls()
            
            # Actualizar pruebas estadísticas
            self.stats_tests_widget.update_columns(self.current_data.columns.tolist())
            
            # Actualizar generador de informes
            self.report_generator.update_sections()
    
    def update_stats(self):
        """Calcula y muestra estadísticas descriptivas avanzadas"""
        if self.current_data is None:
            return
            
        # Estadísticas descriptivas básicas
        stats = self.current_data.describe(include='all').T
        stats['missing'] = self.current_data.isna().sum()
        stats['unique'] = self.current_data.nunique()
        stats['dtype'] = self.current_data.dtypes
        
        # Calcular moda (puede haber múltiples modas)
        modes = self.current_data.mode().iloc[0]
        stats['mode'] = modes
        
        # Calcular asimetría y curtosis para columnas numéricas
        numeric_cols = self.current_data.select_dtypes(include=[np.number]).columns
        stats['skewness'] = np.nan
        stats['kurtosis'] = np.nan
        
        for col in numeric_cols:
            stats.loc[col, 'skewness'] = self.current_data[col].skew()
            stats.loc[col, 'kurtosis'] = self.current_data[col].kurtosis()
        
        self.stats_model.set_dataframe(stats)
        
        # Actualizar opciones para estadísticas por grupo
        self.group_combo.clear()
        categorical_cols = self.current_data.select_dtypes(include=['object', 'category']).columns
        self.group_combo.addItems([""] + list(categorical_cols))
    
    def update_group_stats(self):
        """Actualiza las estadísticas por grupo seleccionado"""
        group_col = self.group_combo.currentText()
        if not group_col or self.current_data is None:
            return
            
        numeric_cols = self.current_data.select_dtypes(include=[np.number]).columns
        if len(numeric_cols) == 0:
            return
            
        # Calcular estadísticas por grupo
        group_stats = self.current_data.groupby(group_col)[numeric_cols].agg(['mean', 'median', 'std', 'count'])
        
        # Aplanar el MultiIndex para mejor visualización
        group_stats.columns = ['_'.join(col).strip() for col in group_stats.columns.values]
        group_stats = group_stats.T.reset_index()
        
        self.group_stats_model.set_dataframe(group_stats)
    
    def update_pivot_controls(self):
        """Actualiza los controles para tablas dinámicas"""
        if self.current_data is None:
            return
            
        columns = [""] + list(self.current_data.columns)
        
        self.pivot_row_combo.clear()
        self.pivot_row_combo.addItems(columns)
        
        self.pivot_col_combo.clear()
        self.pivot_col_combo.addItems(columns)
        
        numeric_cols = [""] + list(self.current_data.select_dtypes(include=np.number).columns)
        self.pivot_val_combo.clear()
        self.pivot_val_combo.addItems(numeric_cols)
    
    def update_chart_controls(self):
        """Actualiza los controles para gráficos"""
        if self.current_data is None:
            return
            
        columns = [""] + list(self.current_data.columns)
        numeric_cols = [""] + list(self.current_data.select_dtypes(include=np.number).columns)
        cat_cols = [""] + list(self.current_data.select_dtypes(include=['object', 'category']).columns)
        
        self.chart_x_combo.clear()
        self.chart_x_combo.addItems(columns)
        
        self.chart_y_combo.clear()
        self.chart_y_combo.addItems(numeric_cols)
        
        self.chart_hue_combo.clear()
        self.chart_hue_combo.addItem("Ninguno", "")
        self.chart_hue_combo.addItems(cat_cols)
    
    def build_advanced_pivot(self):
        """Construye una tabla dinámica avanzada"""
        if self.current_data is None:
            return
            
        rows = self.pivot_row_combo.currentText() or None
        cols = self.pivot_col_combo.currentText() or None
        values = self.pivot_val_combo.currentText() or None
        agg_func = self.pivot_agg_combo.currentText()
        
        if not values:
            return
            
        try:
            pivot_table = pd.pivot_table(
                self.current_data,
                index=rows,
                columns=cols,
                values=values,
                aggfunc=agg_func,
                fill_value=0,
                margins=True,
                margins_name="Total"
            )
            
            self.pivot_model.set_dataframe(pivot_table)
            self.status_bar.showMessage("Tabla dinámica generada", 3000)
            
        except Exception as e:
            self.status_bar.showMessage(f"Error al crear tabla dinámica: {e}", 5000)
    
    def export_pivot(self):
        """Exporta la tabla dinámica actual"""
        if not hasattr(self, 'pivot_model') or self.pivot_model.rowCount() == 0:
            return
            
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Exportar tabla dinámica",
            "",
            "Excel (*.xlsx);;CSV (*.csv)",
            options=options
        )
        
        if not file_path:
            return
        
        try:
            if file_path.endswith('.xlsx'):
                self.pivot_model._filtered_data.to_excel(file_path)
            elif file_path.endswith('.csv'):
                self.pivot_model._filtered_data.to_csv(file_path)
            
            self.status_bar.showMessage("Tabla dinámica exportada", 3000)
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo exportar la tabla:\n{str(e)}")
    
    def generate_advanced_chart(self):
        """Genera el gráfico avanzado seleccionado"""
        if self.current_data is None:
            return
            
        chart_type = self.chart_type_combo.currentText()
        x_col = self.chart_x_combo.currentText()
        y_col = self.chart_y_combo.currentText() if self.chart_y_combo.currentIndex() > 0 else None
        hue_col = self.chart_hue_combo.currentText() if self.chart_hue_combo.currentIndex() > 0 else None
        
        if not x_col:
            return
            
        try:
            if chart_type == "Histograma":
                self.visualizer.plot_histogram(self.current_data, x_col, kde=True)
            elif chart_type == "Dispersión" and y_col:
                self.visualizer.plot_scatter(self.current_data, x_col, y_col, hue=hue_col)
            elif chart_type == "Barras" and y_col:
                self.visualizer.plot_bar(self.current_data, x_col, y_col, hue=hue_col)
            elif chart_type == "Caja" and y_col:
                self.visualizer.plot_box(self.current_data, x_col, y_col, hue=hue_col)
            elif chart_type == "Matriz de Correlación":
                self.visualizer.plot_correlation_matrix(self.current_data)
            elif chart_type == "Gráfico de Violín" and y_col:
                self.figure.clear()
                ax = self.figure.add_subplot(111)
                sns.violinplot(data=self.current_data, x=x_col, y=y_col, hue=hue_col, ax=ax)
                self.canvas.draw()
            elif chart_type == "Gráfico de Densidad":
                self.figure.clear()
                ax = self.figure.add_subplot(111)
                sns.kdeplot(data=self.current_data, x=x_col, hue=hue_col, ax=ax)
                self.canvas.draw()
            
            self.status_bar.showMessage(f"Gráfico {chart_type} generado", 3000)
            
        except Exception as e:
            self.status_bar.showMessage(f"Error al generar gráfico: {e}", 5000)
            QMessageBox.critical(self, "Error", f"No se pudo generar el gráfico:\n{str(e)}")
    
    def export_data(self):
        """Exporta los datos actuales (filtrados o no)"""
        if self.current_data is None:
            return
            
        options = QFileDialog.Options()
        file_path, _ = QFileDialog.getSaveFileName(
            self,
            "Exportar datos",
            "",
            "Excel (*.xlsx);;CSV (*.csv);;JSON (*.json);;HTML (*.html)",
            options=options
        )
        
        if not file_path:
            return
        
        try:
            data_to_export = self.data_model._filtered_data if hasattr(self.data_model, '_filtered_data') else self.current_data
            
            if file_path.endswith('.xlsx'):
                data_to_export.to_excel(file_path, index=False)
            elif file_path.endswith('.csv'):
                data_to_export.to_csv(file_path, index=False)
            elif file_path.endswith('.json'):
                data_to_export.to_json(file_path, orient='records')
            elif file_path.endswith('.html'):
                data_to_export.to_html(file_path, index=False)
            
            self.status_bar.showMessage("Datos exportados correctamente", 3000)
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo exportar los datos:\n{str(e)}")
    
    def run_quick_analysis(self):
        """Ejecuta un análisis rápido de los datos"""
        if self.current_data is None:
            return
            
        try:
            # Generar resumen rápido
            num_cols = self.current_data.select_dtypes(include=[np.number]).columns
            cat_cols = self.current_data.select_dtypes(include=['object', 'category']).columns
            
            report = "=== Análisis Rápido ===\n\n"
            report += f"Filas: {len(self.current_data)}\n"
            report += f"Columnas: {len(self.current_data.columns)}\n"
            report += f"Columnas numéricas: {len(num_cols)}\n"
            report += f"Columnas categóricas: {len(cat_cols)}\n"
            report += f"Valores faltantes: {self.current_data.isna().sum().sum()}\n"
            report += f"Valores duplicados: {self.current_data.duplicated().sum()}\n\n"
            
            if len(num_cols) > 0:
                report += "=== Correlaciones Numéricas ===\n"
                corr_matrix = self.current_data[num_cols].corr()
                for i, col1 in enumerate(num_cols):
                    for col2 in num_cols[i+1:]:
                        corr = corr_matrix.loc[col1, col2]
                        if abs(corr) > 0.7:
                            report += f"{col1} y {col2}: {corr:.2f}\n"
                report += "\n"
            
            if len(cat_cols) > 0:
                report += "=== Valores Únicos ===\n"
                for col in cat_cols:
                    unique_count = self.current_data[col].nunique()
                    if unique_count < 10:
                        report += f"{col}: {unique_count} valores ({', '.join(map(str, self.current_data[col].unique()[:5]))}"
                        if unique_count > 5:
                            report += "..."
                        report += "\n"
                    else:
                        report += f"{col}: {unique_count} valores\n"
            
            # Mostrar reporte en un diálogo
            msg_box = QMessageBox(self)
            msg_box.setWindowTitle("Análisis Rápido")
            msg_box.setText(report)
            msg_box.setDetailedText(str(self.current_data.describe(include='all')))
            msg_box.exec_()
            
        except Exception as e:
            QMessageBox.critical(self, "Error", f"No se pudo completar el análisis:\n{str(e)}")
    
    def show_about(self):
        """Muestra el diálogo 'Acerca de'"""
        about_text = """
        <h1>DataAnalyzer Pro</h1>
        <p>Versión 2.0</p>
        <p>Aplicación avanzada para análisis de datos</p>
        <p>Funcionalidades:</p>
        <ul>
            <li>Carga de archivos (XLSX, CSV, TXT)</li>
            <li>Filtrado avanzado de datos</li>
            <li>Estadísticas descriptivas</li>
            <li>Tablas dinámicas</li>
            <li>Visualización con múltiples gráficos</li>
            <li>Pruebas estadísticas</li>
            <li>Generación de informes profesionales</li>
            <li>Exportación de resultados</li>
        </ul>
        <p>© 2023 - Todos los derechos reservados</p>
        """
        
        QMessageBox.about(self, "Acerca de DataAnalyzer Pro", about_text)

def main():
    """Función principal"""
    app = QApplication(sys.argv)
    app.setStyle('Fusion')  # Mejor aspecto visual
    
    # Configuración inicial
    app.setApplicationName("DataAnalyzer Pro")
    app.setApplicationVersion("2.0")
    
    window = DataAnalyzerPro()
    window.showMaximized()
    
    sys.exit(app.exec_())

if __name__ == "__main__":
    main()